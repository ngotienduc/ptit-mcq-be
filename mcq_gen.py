from llama_index.core.agent import ReActAgent
from llama_index.llms.openai import OpenAI
from llama_index.core.tools import QueryEngineTool, ToolMetadata
from llama_index.core import VectorStoreIndex, Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Settings
from llama_index.core import PromptTemplate
from PyPDF2 import PdfReader
from dotenv import load_dotenv
from docx import Document as DocxDocument
import re

load_dotenv()

llm = OpenAI(model="gpt-3.5-turbo-0125")
text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
Settings.text_splitter = text_splitter

PROMPT_TEMPLATE1 = (
    "Dưới đây là một chủ đề về một môn học."
    "\n -----------------------------\n"
    "{context_str}"
    "\n -----------------------------\n"
    "Bạn là một chuyên gia câu hỏi trắc nghiệm, hãy sinh ra câu hỏi trắc nghiệm gồm 4 đáp án dựa trên chủ đề đưa vào và chỉ ra đáp án đúng."
    "Tạo câu hỏi về chủ đề là:  {query_str}"
)
QA_PROMPT1 = PromptTemplate(PROMPT_TEMPLATE1)

PROMPT_TEMPLATE2 = (
    "Đầu vào là 1 câu hỏi trắc nghiệm."
    "\n -----------------------------\n"
    "{context_str}"
    "\n -----------------------------\n"
    "Bạn là một chuyên gia về câu hỏi trắc nghiệm, hãy kiểm tra lại độ chính xác của câu hỏi và chỉnh sửa lại chúng tốt hơn."
    "Đầu ra là 1 lời đánh giá và một câu hỏi trắc nghiệm. Hãy đánh giá và cập nhật câu hỏi:  {query_str}."
    "Đảm bảo định dạng phản hồi là 1 lời đánh giá và 1 câu hỏi trắc nghiệm có 4 đáp án lựa chọn, sau đó chỉ rõ đáp án đúng."
)
QA_PROMPT2 = PromptTemplate(PROMPT_TEMPLATE2)

react_system_header_str = """\
Bạn được thiết kế để phục vụ một chức năng duy nhất: tạo ra các câu hỏi trắc nghiệm.\
Đáp án cuối cùng bạn cần đưa ra là một câu hỏi trắc nghiệm với bốn lựa chọn và chỉ định câu trả lời đúng.\

## Tools
Bạn có quyền truy cập vào nhiều công cụ khác nhau.
Bạn có trách nhiệm sử dụng các công cụ theo trình tự bất kỳ mà bạn cho là phù hợp để hoàn thành nhiệm vụ trước mắt.
Điều này có thể yêu cầu chia nhiệm vụ thành các nhiệm vụ phụ và sử dụng các công cụ khác nhau để hoàn thành từng nhiệm vụ phụ.

Bạn có quyền truy cập vào các công cụ sau:
{tool_desc}

## Output Format
Để tạo hoặc đánh giá các câu hỏi trắc nghiệm, vui lòng sử dụng định dạng sau.

```
Thought: Tôi cần sử dụng một công cụ để giúp tôi tạo hoặc đánh giá các câu hỏi trắc nghiệm.
Action: tên công cụ (một trong {tool_names}) nếu sử dụng một công cụ.
Action Input: đầu vào của công cụ, ở định dạng JSON biểu thị kwargs (ví dụ: {{"input": "hello world", "num_beams": 5}})
```

Hãy LUÔN bắt đầu bằng một Suy nghĩ.

Vui lòng sử dụng định dạng JSON hợp lệ cho Đầu vào hành động. KHÔNG làm điều này {{'input': 'hello world', 'num_beams': 5}}.

Nếu định dạng này được sử dụng, người dùng sẽ phản hồi theo định dạng sau:

```
Observation: công cụ phản hồi

```

Bạn nên tiếp tục lặp lại định dạng trên cho đến khi có đủ thông tin để đáp ứng yêu cầu
mà không cần sử dụng thêm bất kỳ công cụ nào. Lúc đó bạn PHẢI trả lời ở một trong hai định dạng sau:

```
Thought: Tôi có thể đưa ra câu hỏi trắc nghiệm, mà không cần sử dụng thêm bất kỳ công cụ nào.
Answer: Câu trả lời phải là một câu hỏi trắc nghiệm với bốn lựa chọn và chỉ định câu trả lời đúng.  [câu trả lời ở đây]
```

```
Thought: Tôi không thể trả lời yêu cầu bằng các công cụ được cung cấp.
Answer: Xin lỗi, tôi không thể trả lời yêu cầu của bạn.
```

## Additional Rules
- Câu trả lời PHẢI là một câu hỏi trắc nghiệm với bốn lựa chọn và chỉ định câu trả lời đúng, được viết bằng tiếng việt.
- Câu trả lời PHẢI có định dạng sau: 
Câu hỏi: (Đưa ra câu hỏi)
A. Đáp án 1
B. Đáp án 2
C. Đáp án 3
D. Đáp án 4
Đáp án đúng: (Chỉ ra đáp án đúng)


---
HÃY CHẮC CHẮN CÂU HỎI TRẮC NGHIỆM ĐƯỢC VIẾT BẰNG TIẾNG VIỆT VÀ ĐÚNG ĐỊNH DẠNG TRÊN. 
## Current Conversation
Dưới đây là cuộc trò chuyện hiện tại bao gồm các tin nhắn của con người và trợ lý xen kẽ nhau.

"""

#check xem văn bản có định dạng list hay không
def is_list_format(text):
    try:
        parsed_list = eval(text)
        return isinstance(parsed_list, list)
    except:
        return False

def extract_list_from_text(text):
    pattern = r"\[(.*?)\]"
    match = re.search(pattern, text)
    if match:
        extracted_list = eval(match.group(0))
        return extracted_list
    return []

data = None
def select_topic(topic, quantity):
    PROMPT_TEMPLATE_SELECT_TOPIC = (
        "Dưới đây là một chủ đề về một môn học."
        "\n -----------------------------\n"
        "{context_str}"
        "\n -----------------------------\n"
        "Bạn là một chuyên gia tạo câu hỏi trắc nghiệm, từ một chủ đề được yêu cầu, hãy tìm ra các nội dung có thể sử dụng để tạo câu hỏi trắc nghiệm."
        "Các nội dung được chọn phải có liên quan đến chủ đề được yêu cầu, được viết khái quát và ngắn gọn."
        "Đảm bảo định dạng câu trả lời bao gồm các nội dung được viết theo kiểu list trong python."
        "Ví dụ: ['nội dung 1', 'nội dung 2', 'nội dung 3',...]"
        "Đưa ra danh sách nội dung theo kiểu list trong Python: {query_str}"
    )
    QA_PROMPT_SELECT_TOPIC = PromptTemplate(PROMPT_TEMPLATE_SELECT_TOPIC)

    select_topic_prompt = ''
    query_engine0 = ''
    if topic != '':
        select_topic_prompt ="Hãy chọn " + str(
            quantity) + " nội dung liên quan đến chủ đề \"" + topic + "\" và đưa ra duy nhất một câu trả lời đúng định dạng kiểu list trong python."
    else:
        select_topic_prompt = "Hãy chọn " + str(
            quantity) + " nội dung bất kì trong dữ liệu bạn có và đưa ra duy nhất một câu trả lời đúng định dạng kiểu list trong python."
        
    query_engine0 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT_SELECT_TOPIC,
                                            llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                            max_tokens=-1)
    
    print('select_topic_prompt:', select_topic_prompt)
    response = query_engine0.query(select_topic_prompt)
    print("response: ", response)
    
    # nếu câu trả lời đã ở dạng list thì parse thành list 
    if is_list_format(str(response)):
        subTopics = eval(str(response)) 
        
    # nếu câu trả lời là 1 văn bản chứa list thì trích xuất ra list    
    subTopics = extract_list_from_text(str(response))
    print('subTopics:', subTopics)
    return subTopics

def check(s):
    if s.startswith("Câu") and s.find("A.")!=-1 and s.find("B.")!=-1 and s.find("C.")!=-1 and s.find("D.")!=-1 and s.lower().find("đáp án")!=-1:
        return True
    return False

def read_pdf_file(file):
    print('----------------file-pdf---------------')
    file_content = ''
    reader = PdfReader(file)
    num_pages = len(reader.pages)
    for page_num in range(num_pages):
        page_object = reader.pages[page_num]
        text = page_object.extract_text()
        file_content += " " + text
    return file_content


def read_docx_file(file):
    file_content = ''
    print('----------------file-docx---------------')
    doc = DocxDocument(file)
    for para in doc.paragraphs:
        file_content += " " + para.text
    return file_content


def read_txt_file(file):
    print('----------------file-txt---------------')
    file_content = file.read().decode('utf-8')
    return file_content

def mcqGen(topic, quantity, difficulty, file, inputText, status):
    global data
    if status == 'true':
        file_content = ""
        if file is not None:
            ext_file = file.filename.split('.')[-1]
            if ext_file == 'pdf':
                file_content = read_pdf_file(file)
            elif ext_file == 'docx':
                file_content = read_docx_file(file)
            elif ext_file == 'txt':
                file_content = read_txt_file(file)
            else:
                raise ValueError("Unsupported file type")

        text_splitter = SentenceSplitter(chunk_size=512, chunk_overlap=10)
        content = file_content if file is not None else inputText
        gpt_documents = [Document(text=content)]
        data = VectorStoreIndex.from_documents(documents=gpt_documents, transformations=[text_splitter])

    # GPT
    query_engine1 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT1,
                                         llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.5, max_tokens=512),
                                         max_tokens=-1)
    query_engine2 = data.as_query_engine(similarity_top_k=3, text_qa_template=QA_PROMPT2,
                                         llm=OpenAI(model='gpt-3.5-turbo-0125', temperature=0.1, max_tokens=512),
                                         max_tokens=-1)

    query_engine_tools = [
        QueryEngineTool(
            query_engine=query_engine1,
            metadata=ToolMetadata(
                name="Create",
                description=(
                    "Đầu vào là một yêu cầu. Đầu ra là một câu hỏi trắc nghiệm và có chỉ rõ đáp án đúng."
                    "Tạo câu hỏi trắc nghiệm về chủ đề được yêu cầu."
                    "Sử dụng các công cụ khác để đánh giá câu hỏi."
                ),
            ),
        ),
        QueryEngineTool(
            query_engine=query_engine2,
            metadata=ToolMetadata(
                name="Check1",
                description=(
                    "Đầu vào là một câu hỏi trắc nghiệm. Đầu ra là 1 câu đánh giá và 1 câu hỏi trắc nghiệm. Hãy chỉ rõ câu trả lời đúng."
                    "Tiến hành đánh giá câu hỏi. Giải thích câu trả lời đúng, nếu câu hỏi hoặc câu trả lời sai thì thực hiện chỉnh sửa lại."
                    "Nếu không có câu trả lời đúng thì hãy sửa lại câu trả lời."
                    "Nếu các đáp án tương tự nhau thì hãy sửa lại."
                    "Cải thiện câu hỏi trắc nghiệm."
                    "Kết quả cuối cùng là 1 câu hỏi trắc nghiệm."
                ),
            ),
        ),
        QueryEngineTool(
            query_engine=query_engine2,
            metadata=ToolMetadata(
                name="Check2",
                description=(
                    "Đầu vào là một câu hỏi trắc nghiệm. Đầu ra là 1 câu đánh giá và 1 câu hỏi trắc nghiệm. Hãy chỉ rõ câu trả lời đúng."
                    "Tiến hành đánh giá câu hỏi. Giải thích câu trả lời đúng, nếu câu hỏi hoặc câu trả lời sai thì thực hiện chỉnh sửa lại."
                    "Nếu không có câu trả lời đúng thì hãy sửa lại câu trả lời."
                    "Nếu các đáp án tương tự nhau thì hãy sửa lại."
                    "Cải thiện câu hỏi trắc nghiệm."
                    "Kết quả cuối cùng là 1 câu hỏi trắc nghiệm."
                ),
            ),
        ),
    ]

    agent = ReActAgent.from_tools(
        query_engine_tools,
        llm=llm,
        verbose=True,
    )

    react_system_prompt = PromptTemplate(react_system_header_str)
    agent.update_prompts({"agent_worker:system_prompt": react_system_prompt})
    agent.reset()
    subTopics = select_topic(topic, quantity)

    while len(subTopics) != int(quantity): 
        subTopics = select_topic(topic, quantity )

    mcqs = []
    for subtopic in subTopics:
        kq=""
        while True: 
            prompt = "Tạo 1 câu hỏi trắc nghiệm có nội dung liên quan đến \"" + subtopic.strip(".") + "\""
            if topic != '':
                prompt = prompt + " trong chủ đề \"" + topic + "\""
            if difficulty == "easy":
                prompt = prompt + ". Các câu hỏi có độ khó ở mức thấp, người trả lời có thể dễ dàng tìm ra đáp án đúng"
            if difficulty == "normal":
                prompt = prompt + ". Các câu hỏi có độ khó ở mức trung bình, các đáp án có thể gây nhầm lẫn và yêu cầu người trả lời phải áp dụng một vài bước suy luận"
            if difficulty == "hard":
                prompt = prompt + "Các câu hỏi có độ khó ở mức cao, các đáp án dễ gây ra sự nhầm lẫn, người trả lời phải vận dụng khả năng suy luận và chọn lọc thông tin kỹ càng để chọn ra đáp án đúng"
            prompt = prompt + ". Sau đó sử dụng công cụ kiểm tra lại."
            print("In ra prompt: ")
            print(prompt)
            question = agent.chat(prompt)
            if check(str(question)): 
                kq=question
                break
                
        mcqs.append(str(kq))
    return mcqs



